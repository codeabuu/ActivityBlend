import unittest
from unittest.mock import MagicMock, patch
import docx
from copy import deepcopy
from docx.shared import Pt

# Import the functions from the script
from major.word import create_tables, fill_tables, make_word_doc

class TestScheduleGeneration(unittest.TestCase):

    @patch('major.word.docx.Document')
    def test_create_tables(self, MockDoc):
        mock_doc = MockDoc.return_value
        mock_doc.tables = [MagicMock()]
        
        mock_group_style = MagicMock()
        mock_doc.styles = {'group': mock_group_style}
        create_tables(3, mock_doc)
        
        # Check if tables were created correctly
        self.assertEqual(mock_doc.add_paragraph.call_count, 8)
        # Check if the correct text was added
        for i in range(2, 5):
            mock_doc.add_paragraph.assert_any_call(f"Group {i}")
        
        self.assertEqual(mock_doc.add_page_break.call_count, 3)
        
    def test_fill_tables(self):
        # Create a mock document
        mock_doc = MagicMock()
        
        # Create mock tables with mock cells
        mock_table = MagicMock()
        mock_rows = [MagicMock() for _ in range(3)]
        for row in mock_rows:
            row.cells = [MagicMock() for _ in range(4)]
        mock_table.rows = mock_rows
        mock_doc.tables = [mock_table]

        # Create a matrix to fill the table
        matrix = [
            [["Activity1", "Activity2", "Activity3", "Activity4"],
             ["Activity5", "Activity6", "Activity7", "Activity8"],
             ["Activity9", "Activity10", "Activity11", "Activity12"]]
        ]

        # Run the function
        fill_tables(matrix, mock_doc)

        # Check if cell texts were updated correctly
        for i, row in enumerate(mock_rows):
            for j, cell in enumerate(row.cells):
                self.assertEqual(cell.text, f"Activity{i*4 + j + 1}")

    @patch('your_script.docx.Document')
    @patch('your_script.create_tables')
    @patch('your_script.fill_tables')
    def test_make_word_doc(self, mock_fill_tables, mock_create_tables, MockDoc):
        mock_doc = MockDoc.return_value
        
        # Run the function
        matrix = [
            [["Activity1", "Activity2", "Activity3", "Activity4"],
             ["Activity5", "Activity6", "Activity7", "Activity8"],
             ["Activity9", "Activity10", "Activity11", "Activity12"]]
        ]
        make_word_doc(matrix, "TestWeek")
        
        # Check if Document was opened correctly
        MockDoc.assert_called_with(r"C:\Users\Admin\Desktop\prp\ActivityBlend\major\2024 Template Schedules.docx")
        
        # Check if styles were set correctly
        style = mock_doc.styles['Normal']
        self.assertEqual(style.font.name, 'Arial')
        self.assertEqual(style.font.size, Pt(12))
        
        # Check if create_tables was called with correct arguments
        mock_create_tables.assert_called_once_with(len(matrix)-1, mock_doc)
        
        # Check if the document was saved correctly
        mock_doc.save.assert_called_with("C:\\Users\\Admin\\Desktop\\prp\\ActivityBlend\\major\\Generated Schedules/TestWeek Schedules.docx")

if __name__ == '__main__':
    unittest.main()
