'''
This script is responsible for writing a generated schedule
into a word document.
'''

import docx
from copy import deepcopy
from docx.shared import Pt
import os

GAMES = ["name games", "softball", "basketball", "squash", "ultimate", "hockey", "lacrosse", "american-football", "tennis", "volleyball", "football"]

def create_tables(n, doc):
    '''
    this method creates all of the neccessary tables/schedules for
	each week in a new word document.
    :param doc: word document to add tables to
    :param n: number of tables to add
    '''

    for x in range(n):
        p = doc.add_paragraph("Group " + str(x+3))
        p.style = "group"
        template = doc.tables[0]
        tbl = template._tbl
        new_tbl = deepcopy(tbl)
        paragraph = doc.add_paragraph()
        paragraph._p.addnext(new_tbl)
        doc.add_page_break()

def fill_tables(matrix, doc):
    '''
    fill tables created by create_tables func
    '''
    for group, line in enumerate(doc.tables):
        r = 0
        for row in line.rows:
            col = 0
            update = False
            for cell in row.cells:
                if cell.text == "" or cell.text.lower().strip() == "name games":
                    try:
                        cell.text = matrix[group][r][col]
                        if cell.text == "Name Games":
                            cell.paragraphs[0].style = 'name'
                        else:
                            cell.paragraphs[0].style = 'center'
                        
                        col += 1
                        update = True
                    except:
                        break
            if update:
                r += 1

def make_word_doc(matrix, file_name='Week 1'):
    '''
    create the word doc that contains our new schedule.
    '''

    doc = docx.Document("2024 Template Schedules.docx")
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(13)

    groups = len(matrix)
    create_tables(groups-1, doc)
    doc.save("Generated Schedules/" + file_name + " Activity Schedules.docx")



matrix = [
    [
        ["volleyball", "softball", "basketball"], ["squash", "basketball", "softball"]
        ],
    [
        ["football", "hockey", "lacrosse"], ["football", "tennis", "volleyball"]
        ]
]

make_word_doc(matrix)

